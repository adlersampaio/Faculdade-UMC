from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_style(doc, text, level=1):
    """Adiciona um heading com estilo"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table_with_data(doc, headers, rows):
    """Adiciona uma tabela formatada"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Cor de fundo do header
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D3D3D3')
        header_cells[i]._element.get_or_add_tcPr().append(shading_elm)

    # Dados
    for row_idx, row in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row):
            row_cells[col_idx].text = cell_text

    return table

# Criando o documento
doc = Document()

# Título
title = doc.add_heading('📚 DOCUMENTAÇÃO DO PROJETO', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Diário de Treino - Análise Completa')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph()

# ============================================
# SEÇÃO 1: CONFIGURAÇÕES E DEPENDÊNCIAS
# ============================================
doc.add_heading('🏗️ CONFIGURAÇÕES E DEPENDÊNCIAS', level=1)

doc.add_heading('1. pom.xml - Configurações Maven', level=2)
doc.add_paragraph('Define o projeto como Spring Boot 4.0.6 com Java 25')
doc.add_paragraph('Dependências principais:', style='List Bullet')
doc.add_paragraph('spring-boot-starter-data-jpa → Persistência de dados (banco de dados)', style='List Bullet 2')
doc.add_paragraph('spring-boot-starter-thymeleaf → Template engine (HTML dinâmico)', style='List Bullet 2')
doc.add_paragraph('spring-boot-starter-webmvc → Controllers e rotas HTTP', style='List Bullet 2')
doc.add_paragraph('mysql-connector-j → Driver para conectar no MySQL', style='List Bullet 2')
doc.add_paragraph('Dependências de teste', style='List Bullet 2')

doc.add_heading('2. application.properties - Configuração da Aplicação', level=2)
doc.add_paragraph('Nome da app: diariotreino')
doc.add_paragraph('Conexão MySQL:', style='List Bullet')
doc.add_paragraph('URL: localhost:3306/diario_treino', style='List Bullet 2')
doc.add_paragraph('Usuário: root (sem senha por enquanto)', style='List Bullet 2')
doc.add_paragraph('spring.jpa.hibernate.ddl-auto=update → Cria e atualiza tabelas automaticamente no banco', style='List Bullet')

doc.add_page_break()

# ============================================
# SEÇÃO 2: CAMADA DE MODELOS
# ============================================
doc.add_heading('🗂️ CAMADA DE MODELOS (Banco de Dados)', level=1)

doc.add_heading('3. Usuario.java - Modelo de Usuário', level=2)
doc.add_paragraph('Atributos:', style='List Bullet')
doc.add_paragraph('id (gerado automaticamente)', style='List Bullet 2')
doc.add_paragraph('nome (string)', style='List Bullet 2')
doc.add_paragraph('email (string)', style='List Bullet 2')
doc.add_paragraph('senha (string)', style='List Bullet 2')
doc.add_paragraph('perfil ("ADMIN" ou "USER")', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('É uma entidade JPA (tabela no banco: usuario)', style='Normal')
doc.add_paragraph('Relacionamento: Um usuário pode ter MUITAS fichas de treino', style='Normal')
doc.add_paragraph('Tem construtor vazio e construtor com parâmetros', style='Normal')

doc.add_heading('4. FichaTreino.java - Modelo de Treino', level=2)
doc.add_paragraph('Atributos:', style='List Bullet')
doc.add_paragraph('id (gerado automaticamente)', style='List Bullet 2')
doc.add_paragraph('nomeTreino (ex: "Treino A - Peito")', style='List Bullet 2')
doc.add_paragraph('dataCriacao (data automática = hoje)', style='List Bullet 2')
doc.add_paragraph('usuario (quem criou a ficha - CHAVE ESTRANGEIRA)', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('É uma entidade JPA (tabela no banco: ficha_treino)', style='Normal')
doc.add_paragraph('Relacionamento: Muitas fichas pertencem a UM usuário (@ManyToOne)', style='Normal')
doc.add_paragraph('Data de criação é preenchida automaticamente na inicialização', style='Normal')

doc.add_page_break()

# ============================================
# SEÇÃO 3: CAMADA DE REPOSITÓRIOS
# ============================================
doc.add_heading('💾 CAMADA DE REPOSITÓRIOS (Comunicação com Banco)', level=1)

doc.add_heading('5. UsuarioRepository.java', level=2)
doc.add_paragraph('Interface que herda de JpaRepository', style='List Bullet')
doc.add_paragraph('Método customizado:', style='List Bullet')
doc.add_paragraph('findByEmailAndSenha() → Busca usuário pelo email e senha', style='List Bullet 2')

doc.add_heading('6. FichaTreinoRepository.java', level=2)
doc.add_paragraph('Interface que herda de JpaRepository', style='List Bullet')
doc.add_paragraph('Método customizado:', style='List Bullet')
doc.add_paragraph('findByUsuarioId() → Busca todas as fichas de um usuário específico', style='List Bullet 2')

doc.add_page_break()

# ============================================
# SEÇÃO 4: CAMADA DE CONTROLADORES
# ============================================
doc.add_heading('🎮 CAMADA DE CONTROLADORES (Rotas/Endpoints)', level=1)

doc.add_heading('7. LoginController.java - Autenticação de Usuários', level=2)

# Tabela de rotas
add_table_with_data(doc,
    ['Rota', 'Tipo', 'O que faz'],
    [
        ['/', 'GET', 'Mostra a página de login (index.html)'],
        ['/logar', 'POST', 'Recebe email+senha, verifica no banco, cria sessão'],
        ['/painel', 'GET', 'Mostra painel do usuário (pode redirecioná-lo para login se não logado)'],
        ['/sair', 'GET', 'Faz logout (apaga a sessão)']
    ]
)

doc.add_paragraph()
doc.add_paragraph('Lógica de Segurança: Verifica se o usuário está logado via session.getAttribute("usuarioLogado")', style='Normal')

doc.add_heading('8. FichaTreinoController.java - CRUD de Fichas (Usuários Comuns)', level=2)

add_table_with_data(doc,
    ['Rota', 'Tipo', 'O que faz'],
    [
        ['/nova-ficha', 'GET', 'Abre formulário para criar nova ficha'],
        ['/salvar-ficha', 'POST', 'Salva a ficha no banco com o usuário logado']
    ]
)

doc.add_paragraph()
doc.add_paragraph('Lógica: Pega o usuário da sessão e associa a ficha a ele', style='Normal')

doc.add_heading('9. AdminController.java - CRUD de Fichas (Admin Only)', level=2)

add_table_with_data(doc,
    ['Rota', 'Tipo', 'O que faz'],
    [
        ['/admin/painel', 'GET', 'Mostra TODAS as fichas de todos os alunos (painel gerencial)'],
        ['/admin/nova-ficha', 'GET', 'Abre formulário para criar ficha para algum aluno'],
        ['/admin/salvar-ficha', 'POST', 'Salva a ficha associada a um aluno escolhido'],
        ['/admin/editar-ficha/{id}', 'GET', 'Abre tela de edição de uma ficha'],
        ['/admin/atualizar-ficha', 'POST', 'Salva as alterações na ficha'],
        ['/admin/excluir-ficha/{id}', 'GET', 'Deleta a ficha']
    ]
)

doc.add_paragraph()
doc.add_paragraph('Segurança: Todos os endpoints verificam se o usuário é ADMIN', style='List Bullet')
doc.add_paragraph('Só admins podem gerenciar fichas de TODOS os alunos', style='List Bullet')

doc.add_page_break()

doc.add_heading('10. DiariotreinoApplication.java - Classe Principal', level=2)
doc.add_paragraph('Anotação: @SpringBootApplication - Marca como aplicação Spring Boot', style='List Bullet')
doc.add_paragraph('Anotação: @Bean CommandLineRunner - Executa automaticamente ao iniciar', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Lógica:', style='List Bullet')
doc.add_paragraph('Verifica se existe admin com email "admin@diario.com" e senha "123456"', style='List Bullet 2')
doc.add_paragraph('Se NÃO existir, cria um novo admin automaticamente', style='List Bullet 2')
doc.add_paragraph('Se existir, não faz nada', style='List Bullet 2')

doc.add_heading('11. DiariotreinoApplicationTests.java - Testes Unitários', level=2)
doc.add_paragraph('Teste básico que apenas verifica se a aplicação inicia sem erros', style='List Bullet')
doc.add_paragraph('Anotação @SpringBootTest carrega o contexto completo do Spring', style='List Bullet')

doc.add_page_break()

# ============================================
# SEÇÃO 5: CAMADA DE APRESENTAÇÃO
# ============================================
doc.add_heading('🎨 CAMADA DE APRESENTAÇÃO (Frontend)', level=1)

doc.add_heading('12. style.css - Estilos da Aplicação', level=2)
doc.add_paragraph('Seções:', style='List Bullet')
doc.add_paragraph('Tela de Login: Centro da tela, card branco, botão verde', style='List Bullet 2')
doc.add_paragraph('Layout Geral: Menu lateral esquerdo (sidebar) + conteúdo principal à direita', style='List Bullet 2')
doc.add_paragraph('Sidebar: Menu escuro com links de navegação', style='List Bullet 2')
doc.add_paragraph('Botões: Estilos para "Nova Ficha" (azul), "Sair" (vermelho), "Editar" (amarelo)', style='List Bullet 2')
doc.add_paragraph('Tabelas: Estilos para listas de fichas e alunos', style='List Bullet 2')
doc.add_paragraph('Painel: Container branco com sombra e bordas arredondadas', style='List Bullet 2')

doc.add_heading('13. index.html - Página de Login', level=2)
doc.add_paragraph('Componentes:', style='List Bullet')
doc.add_paragraph('Campo E-mail', style='List Bullet 2')
doc.add_paragraph('Campo Senha', style='List Bullet 2')
doc.add_paragraph('Botão "Entrar" (POST para /logar)', style='List Bullet 2')
doc.add_paragraph('Exibe erro se credenciais estiverem erradas', style='List Bullet 2')

doc.add_heading('14. painel.html - Painel do Usuário Comum', level=2)
doc.add_paragraph('Mostra:', style='List Bullet')
doc.add_paragraph('Bem-vindo com nome do usuário (de session.usuarioLogado)', style='List Bullet 2')
doc.add_paragraph('Botão "Sair"', style='List Bullet 2')
doc.add_paragraph('Botão "+ Criar Nova Ficha de Treino"', style='List Bullet 2')
doc.add_paragraph('Mensagem: "Você ainda não tem nenhuma ficha cadastrada"', style='List Bullet 2')

doc.add_heading('15. admin-painel.html - Painel Gerencial (Admin)', level=2)
doc.add_paragraph('Sidebar com:', style='List Bullet')
doc.add_paragraph('Gerenciar Fichas', style='List Bullet 2')
doc.add_paragraph('Gerenciar Alunos', style='List Bullet 2')
doc.add_paragraph('Sair', style='List Bullet 2')

doc.add_paragraph()
doc.add_paragraph('Conteúdo:', style='List Bullet')
doc.add_paragraph('Titulo "Painel Gerencial - Fichas de Treino"', style='List Bullet 2')
doc.add_paragraph('Botão "+ Nova Ficha de Treino"', style='List Bullet 2')
doc.add_paragraph('Tabela com: Aluno, Treino, Criado em, Ações', style='List Bullet 2')

doc.add_page_break()

doc.add_heading('16. admin-nova-ficha.html - Criar Ficha (Admin)', level=2)
doc.add_paragraph('Formulário com:', style='List Bullet')
doc.add_paragraph('Dropdown para SELECIONAR O ALUNO', style='List Bullet 2')
doc.add_paragraph('Campo "Nome do Treino"', style='List Bullet 2')
doc.add_paragraph('Botão "Salvar Ficha"', style='List Bullet 2')
doc.add_paragraph('Botão "Cancelar"', style='List Bullet 2')

doc.add_heading('17. admin-editar-ficha.html - Editar Ficha (Admin)', level=2)
doc.add_paragraph('Formulário com:', style='List Bullet')
doc.add_paragraph('Campo "Nome do Aluno" (desabilitado - só leitura)', style='List Bullet 2')
doc.add_paragraph('Campo "Nome do Treino" (editável)', style='List Bullet 2')
doc.add_paragraph('Botão "Atualizar Ficha" (amarelo)', style='List Bullet 2')
doc.add_paragraph('Botão "Cancelar"', style='List Bullet 2')

doc.add_heading('18. admin-cadastrar-aluno.html - Registrar Novo Aluno (Admin)', level=2)
doc.add_paragraph('Formulário com:', style='List Bullet')
doc.add_paragraph('Campo "Nome Completo"', style='List Bullet 2')
doc.add_paragraph('Campo "E-mail"', style='List Bullet 2')
doc.add_paragraph('Campo "Senha Provisória"', style='List Bullet 2')
doc.add_paragraph('Botão "Salvar Aluno"', style='List Bullet 2')
doc.add_paragraph('Botão "Cancelar"', style='List Bullet 2')

doc.add_heading('19. admin-alunos.html - Listar Alunos (Admin)', level=2)
doc.add_paragraph('Mostra:', style='List Bullet')
doc.add_paragraph('Título "Alunos Cadastrados"', style='List Bullet 2')
doc.add_paragraph('Botão "+ Novo Aluno"', style='List Bullet 2')
doc.add_paragraph('Tabela com: Nome, E-mail', style='List Bullet 2')

doc.add_page_break()

# ============================================
# SEÇÃO 6: FLUXO GERAL
# ============================================
doc.add_heading('🔄 FLUXO GERAL DA APLICAÇÃO', level=1)

doc.add_paragraph('1. Usuário acessa "/" (login)', style='List Number')
doc.add_paragraph('2. Faz login (email + senha)', style='List Number')
doc.add_paragraph('3. LoginController verifica no banco de dados', style='List Number')
doc.add_paragraph('Se correto → entra em /painel (usuário comum)', style='List Bullet')
doc.add_paragraph('Se errado → fica em "/" com erro', style='List Bullet')
doc.add_paragraph('4. USUÁRIO COMUM pode:', style='List Number')
doc.add_paragraph('Clicar "Nova Ficha" → /nova-ficha', style='List Bullet')
doc.add_paragraph('Ver seu painel', style='List Bullet')
doc.add_paragraph('5. ADMIN pode:', style='List Number')
doc.add_paragraph('Acessar /admin/painel → ver TODAS as fichas', style='List Bullet')
doc.add_paragraph('Criar ficha para aluno', style='List Bullet')
doc.add_paragraph('Editar fichas', style='List Bullet')
doc.add_paragraph('Deletar fichas', style='List Bullet')
doc.add_paragraph('Acessar /admin/alunos → ver todos os alunos', style='List Bullet')
doc.add_paragraph('Cadastrar novos alunos', style='List Bullet')

doc.add_page_break()

# ============================================
# SEÇÃO 7: OBSERVAÇÕES IMPORTANTES
# ============================================
doc.add_heading('⚠️ OBSERVAÇÕES IMPORTANTES', level=1)

doc.add_heading('❌ Problemas/Não Implementado:', level=2)
doc.add_paragraph('Senhas NÃO estão criptografadas (SEGURANÇA EM RISCO!)', style='List Bullet')
doc.add_paragraph('Validações de input fraca', style='List Bullet')
doc.add_paragraph('Controllers para admin ainda não têm todas as rotas implementadas', style='List Bullet')
doc.add_paragraph('Tela de painel do usuário comum não lista suas fichas', style='List Bullet')
doc.add_paragraph('Alguns endpoints faltam as rotas corretas', style='List Bullet')

doc.add_heading('✅ O que está funcionando:', level=2)
doc.add_paragraph('Login/Logout', style='List Bullet')
doc.add_paragraph('Criação automática do admin', style='List Bullet')
doc.add_paragraph('Controle de sessão', style='List Bullet')
doc.add_paragraph('Banco de dados com relacionamentos', style='List Bullet')
doc.add_paragraph('Interface com Thymeleaf', style='List Bullet')

# Salvando o documento
output_path = r'C:\Users\adler\OneDrive\Área de Trabalho\Pessoal\diariotreino\DOCUMENTACAO_PROJETO_DIARIO_TREINO.docx'
doc.save(output_path)
print(f"✅ Documento criado com sucesso em: {output_path}")

